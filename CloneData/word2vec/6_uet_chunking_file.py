import re
import openpyxl
import pandas as pd
from bs4 import BeautifulSoup
from sqlalchemy import create_engine, text
from sqlalchemy.engine import Result
from transformers import AutoTokenizer
import pdfplumber
from docx import Document

TOKENIZER_NAME = "vinai/phobert-large"
MAX_SEQ_LEN    = 250  # 510 + [CLS] + [SEP] = 256

tokenizer = AutoTokenizer.from_pretrained(TOKENIZER_NAME)
engine = create_engine("mysql+pymysql://root:root@localhost/chatbot", future=True)

SQL_INSERT = text("""
    INSERT INTO uet_chunking (id_content, main_title, sub_title, content)
    VALUES (:id_content, :main_title, :sub_title, :content);
""")

SQL_UPDATE = text("""
    UPDATE uet_content SET chunking = 1 WHERE id = :id_content;
""")

def excel():
    def flatten_cols(col_tuple):
        # gom các level của header vào một tên cột đơn
        parts = []
        for v in col_tuple:
            if isinstance(v, str) and v.strip() and not v.lower().startswith("unnamed"):
                # bỏ newline và trim
                parts.append(v.replace("\n", " ").strip())
        return re.sub(r"\s+", " ", " ".join(parts)).lower()
    def excel_to_text(path):
        wb = openpyxl.load_workbook(path, data_only=True)

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            max_col = ws.max_column

            # --- 1) TÌM TITLE ROWS qua horizontal merges ---
            merges = [m for m in ws.merged_cells.ranges if m.min_row == m.max_row]
            row_to_width = {}
            for m in merges:
                w = m.max_col - m.min_col + 1
                row_to_width.setdefault(m.min_row, []).append(w)
            if row_to_width:
                max_w = max(max(ws) for ws in row_to_width.values())
                title_rows = [r for r, widths in row_to_width.items() if max(widths) == max_w]
                title_parts = []
                for r in title_rows:
                    for m in merges:
                        if m.min_row == r and (m.max_col - m.min_col + 1) == max_w:
                            v = ws.cell(row=m.min_row, column=m.min_col).value
                            if v: title_parts.append(str(v).strip())
                title = " ".join(title_parts)
            else:
                title_rows = []
                title = ""
            print(f"\n=== Sheet: {sheet_name} ===")
            print("Title:", title, "\n")

            # --- 2) TÌM HEADER_START: hàng đầu tiên sau title_rows có dữ liệu ---
            start = max(title_rows) + 1 if title_rows else 1
            header_start = None
            for i in range(start, ws.max_row + 1):
                rowvals = [ws.cell(row=i, column=c).value for c in range(1, max_col+1)]
                if any(v not in (None, "") for v in rowvals):
                    header_start = i
                    break
            if header_start is None:
                continue

            # --- 3) TÌM HEADER_END qua vertical merges ---
            header_end = header_start
            for m in ws.merged_cells.ranges:
                # nếu merge range “vắt qua” header_start
                if m.min_row <= header_start < m.max_row:
                    header_end = max(header_end, m.max_row)

            # --- 4) ĐỌC pandas với tất cả các hàng header ---
            header_rows = list(range(header_start, header_end + 1))
            header_idx  = [r - 1 for r in header_rows]  # pandas dùng 0-based
            df = pd.read_excel(
                path,
                sheet_name=sheet_name,
                header=header_idx,
                dtype=str
            )
            df.columns = [flatten_cols(col) for col in df.columns]

            # --- 5) IN từng dòng data ---
            data = []
            for _, row in df.iterrows():
                # dừng khi STT (cột đầu) trống
                if pd.isna(row.iloc[0]):
                    break
                parts = []
                for col, val in row.items():
                    if pd.notna(val):
                        parts.append(f"{col} {str(val).strip()}")
                data.append(" ".join(parts))
        return [title, data]

    SQL_DATA = text("""
        SELECT id, paragraph
        FROM uet_content
        WHERE type = 'file' and chunking = 0
        AND (
            LOWER(paragraph) LIKE '%.xlsx' 
        );
    """)

    SQL_TITLE = text("""
        SELECT paragraph FROM uet_content WHERE id_url = (SELECT id_parents from uet_url WHERE id = (SELECT id_url from uet_content WHERE id = :id));
    """)
    with engine.connect() as conn:
        rows = conn.execute(SQL_DATA).all()
    for id_content, html in rows:
        sub_title, data = excel_to_text(html)
        with engine.connect() as conn:
            result: Result = conn.execute(SQL_TITLE, {"id": id_content})
            row = result.mappings().first()
            soup = BeautifulSoup(row.get("paragraph"), 'html.parser')
            main_title = soup.title and soup.title.string or ""
        with engine.begin() as tx:
            for d in data:
                tx.execute(
                    SQL_INSERT,
                    {
                        "id_content": id_content,
                        "main_title": main_title,
                        "sub_title": sub_title,
                        "content": d,
                    },
                )
            tx.execute(SQL_UPDATE, {"id_content": id_content})
            tx.commit()
            print(id_content)

def pdf():
    def pdf_to_terminal(pdf_path):
        def chunk_text(text: str, chunk_size: int = 250, overlap: int = 25):
            """
            Chia text thành các chunk có độ dài chunk_size token,
            mỗi chunk chồng lắp overlap token với chunk trước.
            """
            # Mã hoá text thành danh sách token ids (không thêm special tokens)
            token_ids = tokenizer.encode(text, add_special_tokens=False)
            
            chunks = []
            step = chunk_size - overlap
            for start in range(0, len(token_ids), step):
                end = start + chunk_size
                chunk_ids = token_ids[start:end]
                # Giải mã lại thành chuỗi văn bản
                chunk_text = tokenizer.decode(chunk_ids, clean_up_tokenization_spaces=True)
                chunks.append(chunk_text)
                if end >= len(token_ids):
                    break
            return chunks
        
        result = []
        with pdfplumber.open(pdf_path) as pdf:
            total_pages = len(pdf.pages)
            if(total_pages > 10):
                return []
            for page_no, page in enumerate(pdf.pages, start=1):
                # 1) Thử trích bảng
                tables = page.extract_tables({
                    "vertical_strategy": "lines",
                    "horizontal_strategy": "lines"
                })
                
                # 2) Kiểm tra xem có bảng hợp lệ không (ít nhất 2 cột header và ≥2 hàng)
                has_table = False
                if tables:
                    tbl = tables[0]
                    if len(tbl) > 1 and len(tbl[0]) > 1:
                        has_table = True
                
                if has_table:
                    # 3) Chuyển bảng thành DataFrame và in từng dòng flatten
                    df = pd.DataFrame(tables[0][1:], columns=tables[0][0])
                    for row in df.itertuples(index=False, name=None):
                        parts = []
                        for col_name, val in zip(df.columns, row):
                            if val is not None and str(val).strip():
                                parts.append(f"{col_name and col_name.lower() or ''} {str(val).strip()}")
                        result.append(" ".join(parts))
                else:
                    # 4) Fallback: in nguyên văn trang đó
                    text = page.extract_text()
                    if text:
                        result.extend(chunk_text(text.strip()))
        result = [s for s in result if s.strip()]
        return result

    SQL_DATA = text("""
        SELECT id, paragraph
        FROM uet_content
        WHERE type = 'file' and chunking = 0
        AND (
            LOWER(paragraph) LIKE '%.pdf' 
        );
    """)

    SQL_TITLE = text("""
        SELECT paragraph FROM uet_content WHERE id_url = (SELECT id_parents from uet_url WHERE id = (SELECT id_url from uet_content WHERE id = :id));
    """)
    with engine.connect() as conn:
        rows = conn.execute(SQL_DATA).all()
    for id_content, html in rows:
        data = pdf_to_terminal(html)
        with engine.connect() as conn:
            result: Result = conn.execute(SQL_TITLE, {"id": id_content})
            row = result.mappings().first()
            soup = BeautifulSoup(row.get("paragraph"), 'html.parser')
            main_title = soup.title and soup.title.string or ""
        with engine.begin() as tx:
            for d in data:
                tx.execute(
                    SQL_INSERT,
                    {
                        "id_content": id_content,
                        "main_title": main_title,
                        "sub_title": "",
                        "content": d,
                    },
                )
            tx.execute(SQL_UPDATE, {"id_content": id_content})
            tx.commit()
            print(id_content)

def word():
    def word_to_terminal(docx_path: str):
        def chunk_text(text: str, chunk_size: int = 250, overlap: int = 24):
            """
            Chia text thành các chunk có độ dài chunk_size token,
            mỗi chunk chồng lắp overlap token với chunk trước.
            """
            token_ids = tokenizer.encode(text, add_special_tokens=False)
            chunks = []
            step = chunk_size - overlap
            for start in range(0, len(token_ids), step):
                end = start + chunk_size
                chunk_ids = token_ids[start:end]
                chunks.append(tokenizer.decode(chunk_ids, clean_up_tokenization_spaces=True))
                if end >= len(token_ids):
                    break
            return chunks
        result = []
        doc = Document(docx_path)

        # 1) Nếu có bảng thì flatten từng dòng của tất cả các bảng
        if doc.tables:
            for table in doc.tables:
                # Lấy toàn bộ rows dưới dạng list of lists
                rows = [[cell.text for cell in row.cells] for row in table.rows]
                # Kiểm tra tối thiểu phải có header ≥2 cột và ≥2 hàng
                if len(rows) > 1 and len(rows[0]) > 1:
                    df = pd.DataFrame(rows[1:], columns=rows[0])
                    for row in df.itertuples(index=False, name=None):
                        parts = []
                        for col_name, val in zip(df.columns, row):
                            if val is not None and str(val).strip():
                                parts.append(f"{col_name.lower()} {str(val).strip()}")
                        result.append(" ".join(parts))
        else:
            # 2) Nếu không có bảng, fallback: chunk toàn bộ văn bản
            full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            if full_text:
                result.extend(chunk_text(full_text.strip()))

        # Loại bỏ dòng rỗng
        return [s for s in result if s.strip()]

    SQL_DATA = text("""
        SELECT id, paragraph
        FROM uet_content
        WHERE type = 'file' and chunking = 0
        AND (
            LOWER(paragraph) LIKE '%.docx' 
        );
    """)

    SQL_TITLE = text("""
        SELECT paragraph FROM uet_content WHERE id_url = (SELECT id_parents from uet_url WHERE id = (SELECT id_url from uet_content WHERE id = :id));
    """)
    with engine.connect() as conn:
        rows = conn.execute(SQL_DATA).all()
    for id_content, html in rows:
        data = word_to_terminal(html)
        with engine.connect() as conn:
            result: Result = conn.execute(SQL_TITLE, {"id": id_content})
            row = result.mappings().first()
            soup = BeautifulSoup(row.get("paragraph"), 'html.parser')
            main_title = soup.title and soup.title.string or ""
        with engine.begin() as tx:
            for d in data:
                tx.execute(
                    SQL_INSERT,
                    {
                        "id_content": id_content,
                        "main_title": main_title,
                        "sub_title": "",
                        "content": d,
                    },
                )
            tx.execute(SQL_UPDATE, {"id_content": id_content})
            tx.commit()
            print(id_content)

word()