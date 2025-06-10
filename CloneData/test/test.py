#!/usr/bin/env python3
import sys
from docx import Document
import pandas as pd
from transformers import AutoTokenizer

TOKENIZER_NAME = "vinai/phobert-large"
MAX_SEQ_LEN    = 250  # 250 token mỗi chunk
OVERLAP        = 25   # overlap giữa các chunk

tokenizer = AutoTokenizer.from_pretrained(TOKENIZER_NAME)


def word_to_terminal(docx_path: str):
    def chunk_text(text: str, chunk_size: int = MAX_SEQ_LEN, overlap: int = OVERLAP):
        """
        Chia text thành các chunk có độ dài chunk_size token,
        mỗi chunk chồng lắp overlap token với chunk trước.
        """
        token_ids = tokenizer.encode(text, add_special_tokens=False)
        chunks = []
        step = chunk_size - overlap
        for start in range(0, len(token_ids), step):
            end = start + chunk_size
            chunk_ids = token_ids[start:end]
            chunks.append(tokenizer.decode(chunk_ids, clean_up_tokenization_spaces=True))
            if end >= len(token_ids):
                break
        return chunks
    result = []
    doc = Document(docx_path)

    # 1) Nếu có bảng thì flatten từng dòng của tất cả các bảng
    if doc.tables:
        for table in doc.tables:
            # Lấy toàn bộ rows dưới dạng list of lists
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            # Kiểm tra tối thiểu phải có header ≥2 cột và ≥2 hàng
            if len(rows) > 1 and len(rows[0]) > 1:
                df = pd.DataFrame(rows[1:], columns=rows[0])
                for row in df.itertuples(index=False, name=None):
                    parts = []
                    for col_name, val in zip(df.columns, row):
                        if val is not None and str(val).strip():
                            parts.append(f"{col_name.lower()} {str(val).strip()}")
                    result.append(" ".join(parts))
    else:
        # 2) Nếu không có bảng, fallback: chunk toàn bộ văn bản
        full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if full_text:
            result.extend(chunk_text(full_text.strip()))

    # Loại bỏ dòng rỗng
    return [s for s in result if s.strip()]

if __name__ == "__main__":
    for line in word_to_terminal(r"E:\Code\Master\BDT\CloneFile\https___uet.vnu.edu.vn_wp-content_uploads_2024_06_Ph%E1%BB%A5-l%E1%BB%A5c-3.docx.docx"):
        print(line)
